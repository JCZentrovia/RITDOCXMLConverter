"""Enhanced DOCX builder that applies AI formatting patterns."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
    WD_STYLE_TYPE = None

logger = logging.getLogger(__name__)


class FormattedDocxBuilder:
    """Builds formatted DOCX documents from blocks with AI formatting."""
    
    def __init__(self):
        if Document is None:
            raise RuntimeError(
                "python-docx is required. Install with: pip install python-docx"
            )
        
        self.document = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the document."""
        styles = self.document.styles
        
        # Create or get Chapter Heading style
        try:
            chapter_style = styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.size = Pt(18)
            chapter_style.font.bold = True
            # chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chapter_style.paragraph_format.space_before = Pt(24)
            chapter_style.paragraph_format.space_after = Pt(12)
        except:
            pass  # Style might already exist
        
        # Create or get Section Heading style
        try:
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        except:
            pass
    
    def build_from_blocks(
        self,
        blocks: List[Dict],
        output_path: Path,
        metadata: Optional[Dict] = None
    ) -> Path:
        """
        Build formatted DOCX from blocks with AI formatting.
        
        Args:
            blocks: List of blocks (from heuristics + AI enhancement)
            output_path: Path to save DOCX file
            metadata: Optional document metadata
        
        Returns:
            Path to created DOCX file
        """
        logger.info("Building formatted DOCX with %d blocks", len(blocks))
        
        for block in blocks:
            label = block.get("classifier_label") or block.get("label", "para")
            text = (block.get("text") or "").strip()
            
            if not text:
                continue
            
            # Get AI formatting if available
            ai_formatting = block.get("ai_formatting", {})
            
            # Create paragraph based on label
            if label == "book_title":
                self._add_title(text, ai_formatting)
            elif label == "chapter":
                self._add_chapter(text, ai_formatting)
            elif label == "section":
                self._add_section(text, ai_formatting)
            elif label == "toc":
                self._add_toc_entry(text, ai_formatting)
            elif label == "list_item":
                self._add_list_item(text, ai_formatting, block.get("list_type"))
            elif label == "caption":
                self._add_caption(text, ai_formatting)
            elif label == "para":
                self._add_paragraph(text, ai_formatting)
            else:
                # Default to paragraph for unknown labels
                self._add_paragraph(text, ai_formatting)
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))
        logger.info("Saved formatted DOCX to %s", output_path)
        
        return output_path
    
    def _add_title(self, text: str, formatting: Dict):
        """Add book title."""
        para = self.document.add_heading(text, level=0)
        self._apply_formatting(para, formatting)
    
    def _add_chapter(self, text: str, formatting: Dict):
        """Add chapter heading."""
        para = self.document.add_heading(text, level=1)
        self._apply_formatting(para, formatting)
    
    def _add_section(self, text: str, formatting: Dict):
        """Add section heading."""
        para = self.document.add_heading(text, level=2)
        self._apply_formatting(para, formatting)
    
    def _add_toc_entry(self, text: str, formatting: Dict):
        """Add table of contents entry."""
        para = self.document.add_paragraph(text)
        # Don't set a specific style - let AI formatting handle the appearance
        self._apply_formatting(para, formatting)
    
    def _add_list_item(self, text: str, formatting: Dict, list_type: Optional[str]):
        """Add list item."""
        style = 'List Number' if list_type == "ordered" else 'List Bullet'
        para = self.document.add_paragraph(text, style=style)
        self._apply_formatting(para, formatting)
    
    def _add_caption(self, text: str, formatting: Dict):
        """Add caption."""
        para = self.document.add_paragraph(text)
        para.style = 'Caption'
        self._apply_formatting(para, formatting)
    
    def _add_paragraph(self, text: str, formatting: Dict):
        """Add regular paragraph."""
        para = self.document.add_paragraph(text)
        self._apply_formatting(para, formatting)
    
    def _apply_formatting(self, paragraph, formatting: Dict):
        """Apply AI formatting to a paragraph."""
        if not formatting:
            return
        
        # Apply alignment
        # alignment = formatting.get("alignment")
        # if alignment and WD_ALIGN_PARAGRAPH:
        #    align_map = {
        #        "left": WD_ALIGN_PARAGRAPH.LEFT,
        #        "center": WD_ALIGN_PARAGRAPH.CENTER,
        #        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        #        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        #   }
        #    if alignment in align_map:
        #        paragraph.alignment = align_map[alignment]
        
        # Apply spacing
        if "spacing_before" in formatting:
            paragraph.paragraph_format.space_before = Pt(formatting["spacing_before"])
        if "spacing_after" in formatting:
            paragraph.paragraph_format.space_after = Pt(formatting["spacing_after"])
        
        # Apply indentation
        if "indent_first_line" in formatting:
            paragraph.paragraph_format.first_line_indent = Inches(formatting["indent_first_line"])
        if "indent_left" in formatting:
            paragraph.paragraph_format.left_indent = Inches(formatting["indent_left"])
        
        # Apply text formatting (to all runs in paragraph)
        for run in paragraph.runs:
            if formatting.get("bold"):
                run.bold = True
            if formatting.get("italic"):
                run.italic = True
            if formatting.get("underline"):
                run.underline = True
            if "font_size" in formatting:
                run.font.size = Pt(formatting["font_size"])
            if "font_family" in formatting:
                run.font.name = formatting["font_family"]


def create_formatted_docx(
    blocks: List[Dict],
    output_path: Path,
    metadata: Optional[Dict] = None
) -> Path:
    """
    Convenience function to create formatted DOCX from blocks.
    
    Args:
        blocks: List of blocks with AI formatting
        output_path: Path to save DOCX
        metadata: Optional document metadata
    
    Returns:
        Path to created DOCX file
    """
    builder = FormattedDocxBuilder()
    return builder.build_from_blocks(blocks, output_path, metadata)
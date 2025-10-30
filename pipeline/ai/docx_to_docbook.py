from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from lxml import etree

try:  # pragma: no cover - optional dependency
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

logger = logging.getLogger(__name__)


def _normalise_text(text: str) -> str:
    return " ".join(text.split())


def _gather_docx_text(docx_path: Path) -> str:
    if Document is None:
        raise RuntimeError("python-docx is required to inspect DOCX files")

    document = Document(str(docx_path))
    parts = [para.text for para in document.paragraphs if para.text]
    return "\n".join(parts)


def convert_docx_to_docbook(
    docx_path: Path,
    output_path: Optional[Path] = None,
    *,
    expected_text: Optional[str] = None,
    root_name: str = "article",
) -> Path:
    """Convert a formatted DOCX file into a minimal DocBook document.

    The conversion intentionally keeps the mapping simple—section headings become
    ``<title>`` elements and all other paragraphs are wrapped in ``<para>``
    nodes. The routine optionally validates that the textual content exactly
    matches *expected_text* (after whitespace normalisation) to guarantee that
    formatting changes never alter the underlying words.
    """

    if Document is None:
        raise RuntimeError("python-docx is required to convert DOCX files")

    if output_path is None:
        output_path = docx_path.with_suffix(".xml")

    document = Document(str(docx_path))

    if expected_text is not None:
        docx_text = _normalise_text(_gather_docx_text(docx_path))
        if _normalise_text(expected_text) != docx_text:
            raise ValueError(
                "Formatted DOCX text does not match the expected plain text; "
                "aborting DocBook conversion to prevent content drift."
            )

    root = etree.Element(root_name)
    body = etree.SubElement(root, "section")
    current_section = body

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            current_section = etree.SubElement(root, "section")
            title = etree.SubElement(current_section, "title")
            title.text = paragraph.text
        else:
            para_node = etree.SubElement(current_section, "para")
            para_node.text = paragraph.text

    tree = etree.ElementTree(root)
    output_path.write_bytes(
        etree.tostring(
            tree,
            encoding="UTF-8",
            pretty_print=True,
            xml_declaration=True,
        )
    )
    logger.info("Wrote DocBook rendition of %s to %s", docx_path, output_path)
    return output_path

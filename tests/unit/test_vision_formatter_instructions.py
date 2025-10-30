from __future__ import annotations

from pathlib import Path

import pytest

docx_mod = pytest.importorskip("docx")
Document = docx_mod.Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pipeline.ai.vision_formatter import _materialise_docx_from_instructions


def test_materialise_docx_from_instructions(tmp_path: Path):
    plain_text = "Heading\nBody text"
    instructions = {
        "paragraphs": [
            {"line": 1, "style": "Heading 1", "alignment": "center"},
            {"line": 2, "bold": [[0, 4]], "italic": [[5, 9]]},
        ]
    }

    destination = tmp_path / "formatted.docx"
    _materialise_docx_from_instructions(plain_text, instructions, destination)

    document = Document(str(destination))
    assert len(document.paragraphs) == 2

    heading = document.paragraphs[0]
    assert heading.text == "Heading"
    assert heading.style.name == "Heading 1"
    assert heading.alignment == WD_ALIGN_PARAGRAPH.CENTER

    body = document.paragraphs[1]
    assert body.text == "Body text"
    runs = body.runs
    assert [run.text for run in runs] == ["Body", " ", "text"]
    assert runs[0].bold is True
    assert runs[1].bold is None
    assert runs[2].italic is True
